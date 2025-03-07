import time
import pandas as pd
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document  # For Word document creation
from docx.shared import Pt
from docx.oxml import OxmlElement, ns

# **Peterson Institute Research Page URL**
piie_url = "https://www.piie.com/research"

# **Configure Selenium WebDriver**
options = Options()
options.headless = True  # Runs in background
options.add_argument("--no-sandbox")
options.add_argument("--disable-dev-shm-usage")

# Start Selenium WebDriver
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)

# **Open the Peterson Institute webpage**
driver.get(piie_url)
time.sleep(5)  # Allow JavaScript to load

# **Extract the page source and parse with BeautifulSoup**
soup = BeautifulSoup(driver.page_source, "html.parser")

# **Find all article links and titles**
articles = []
for article in soup.find_all("h2", class_="teaser__title"):
    a_tag = article.find("a")
    if a_tag:
        title = a_tag.get_text(strip=True)
        link = urljoin(piie_url, a_tag["href"])  # Ensure full URL

        articles.append({"Think Tank": "Peterson Institute", "Title": title, "URL": link})

# **Close browser**
driver.quit()

# **Save results to CSV**
if articles:
    df = pd.DataFrame(articles)
    df.to_csv("Peterson_Institute_articles.csv", index=False)
    print("✅ Peterson Institute articles scraped and saved to Peterson_Institute_articles.csv")
else:
    print("⚠️ No articles found on Peterson Institute's page.")


# **EXTENSION: Create a Word Document from CSV with Hyperlinked Titles**
# ----------------------------------
word_filename = "Peterson hyperlinks.docx"

# Read data from CSV
df = pd.read_csv("Peterson_Institute_articles.csv")

# Create a Word document
doc = Document()

# Add a title
doc.add_heading("Peterson Institute", level=1)

# Function to add a real clickable hyperlink with custom formatting
def add_hyperlink(paragraph, text, url):
    """
    Adds a real clickable hyperlink to a Word document with specific formatting.
    """
    # Create the hyperlink relationship
    r_id = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(ns.qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Set Calibri font, 13.5 size, underlined, and blue color (Accent 1)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(ns.qn("w:ascii"), "Calibri")
    rFonts.set(ns.qn("w:hAnsi"), "Calibri")

    sz = OxmlElement("w:sz")
    sz.set(ns.qn("w:val"), "27")  # 13.5 pt (size * 2)

    u = OxmlElement("w:u")
    u.set(ns.qn("w:val"), "single")  # Underlined

    color = OxmlElement("w:color")
    color.set(ns.qn("w:val"), "0000FF")  # Blue (Accent 1)

    rPr.append(rFonts)
    rPr.append(sz)
    rPr.append(u)
    rPr.append(color)
    r.append(rPr)

    # Add text
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)

    hyperlink.append(r)
    paragraph._element.append(hyperlink)

# Add articles to Word document
for index, row in df.iterrows():
    title = row["Title"]
    url = row["URL"]

    # Add hyperlink title with custom formatting
    p = doc.add_paragraph()
    add_hyperlink(p, title, url)

    # Add URL below the title
    doc.add_paragraph(url)

# Save the document
doc.save(word_filename)

# **Save the document in the output_files folder**
word_file_path = "output_files/PIIE.docx"
doc.save(word_file_path)

print(f"✅ Word document created: {word_filename}")
# ----------------------------------

