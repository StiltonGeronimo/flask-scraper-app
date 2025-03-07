import requests
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document  # For Word document creation
from docx.shared import Pt
from docx.oxml import OxmlElement, ns

# **Heritage Research Page URL**
heritage_url = "https://www.heritage.org/"

# **Fetch the page using BeautifulSoup**
headers = {"User-Agent": "Mozilla/5.0"}
response = requests.get(heritage_url, headers=headers)

# **Check if the request was successful**
if response.status_code == 200:
    soup = BeautifulSoup(response.text, "html.parser")

    # **Extract article titles and links**
    articles = []
    for article in soup.find_all("h4", class_="view-list--header"):
        title = article.get_text(strip=True)
        
        # Find the parent <a> tag if available for the link
        a_tag = article.find_parent("a")
        link = a_tag["href"] if a_tag and "href" in a_tag.attrs else "No URL Found"

        articles.append({"Think Tank": "Heritage", "Title": title, "URL": link})

    # **Save results to CSV**
    if articles:
        df = pd.DataFrame(articles)
        df.to_csv("Heritage_articles.csv", index=False)
        print("✅ Heritage articles scraped and saved to Heritage_articles.csv")
    else:
        print("⚠️ No articles found on Heritage's page.")

else:
    print(f"❌ Failed to fetch Heritage page: {response.status_code}")

# **EXTENSION: Create a Word Document from CSV with Hyperlinked Titles**
# ----------------------------------
word_filename = "Heritage hyperlinks.docx"

# Read data from CSV
df = pd.read_csv("Heritage_articles.csv")

# Create a Word document
doc = Document()

# Add a title
doc.add_heading("Heritage Foundation", level=1)

# Function to add a real clickable hyperlink with custom formatting
def add_hyperlink(paragraph, text, url):
    """
    Adds a real clickable hyperlink to a Word document with specific formatting.
    """
    # Create the hyperlink relationship
    r_id = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(ns.qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Set Calibri font, 13.5 size, underlined, and blue color (Accent 1)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(ns.qn("w:ascii"), "Calibri")
    rFonts.set(ns.qn("w:hAnsi"), "Calibri")

    sz = OxmlElement("w:sz")
    sz.set(ns.qn("w:val"), "27")  # 13.5 pt (size * 2)

    u = OxmlElement("w:u")
    u.set(ns.qn("w:val"), "single")  # Underlined

    color = OxmlElement("w:color")
    color.set(ns.qn("w:val"), "0000FF")  # Blue (Accent 1)

    rPr.append(rFonts)
    rPr.append(sz)
    rPr.append(u)
    rPr.append(color)
    r.append(rPr)

    # Add text
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)

    hyperlink.append(r)
    paragraph._element.append(hyperlink)

# Add articles to Word document
for index, row in df.iterrows():
    title = row["Title"]
    url = row["URL"]

    # Add hyperlink title with custom formatting
    p = doc.add_paragraph()
    add_hyperlink(p, title, url)

    # Add URL below the title
    doc.add_paragraph(url)

# Save the document
doc.save(word_filename)

# **Save the document in the output_files folder**
word_file_path = "output_files/Heritage Foundation.docx"
doc.save(word_file_path)

print(f"Word document created: {word_filename}")
# ----------------------------------

