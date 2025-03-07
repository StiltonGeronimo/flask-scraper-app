import requests
import pandas as pd
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document  # For Word document creation
from docx.shared import Pt
from docx.oxml import OxmlElement, ns

# **MEI Research Page URL**
mei_url = "https://www.mei.edu/policy-analysis"

# **Fetch the page using BeautifulSoup**
headers = {"User-Agent": "Mozilla/5.0"}
response = requests.get(mei_url, headers=headers)

# **Check if the request was successful**
if response.status_code == 200:
    soup = BeautifulSoup(response.text, "html.parser")

    # **Extract article titles, links, and dates**
    articles = []
    for article in soup.find_all("article", class_="feature feature-1")[:20]:  # Limit to 20 articles
        # Extract title safely
        title_tag = article.find("h4")
        title_link = title_tag.find("a") if title_tag else None
        title = title_link.get_text(strip=True) if title_link else "No Title"

        # Extract article URL
        link = urljoin(mei_url, title_link["href"]) if title_link else "No URL"

        # Extract publication date
        date_tag = article.find("span", class_="feature__date")
        date = date_tag.get_text(strip=True) if date_tag else "No Date"

        # Append extracted data
        if title != "No Title" and link != "No URL":  # Ensure valid articles
            articles.append({"Think Tank": "MEI", "Date": date, "Title": title, "URL": link})

    # **Save results to CSV**
    if articles:
        df = pd.DataFrame(articles)
        df.to_csv("MEI_articles.csv", index=False)
        print("✅ MEI articles scraped and saved to MEI_articles.csv")
    else:
        print("⚠️ No articles found on MEI's page.")

else:
    print(f"❌ Failed to fetch MEI page: {response.status_code}")

# **EXTENSION: Create a Word Document from CSV with Hyperlinked Titles**
# ----------------------------------
word_filename = "MEI hyperlinks.docx"

# Read data from CSV
df = pd.read_csv("MEI_articles.csv")

# Create a Word document
doc = Document()

# Add a title
doc.add_heading("MEI", level=1)

# Function to add a real clickable hyperlink with custom formatting
def add_hyperlink(paragraph, text, url):
    """
    Adds a real clickable hyperlink to a Word document with specific formatting.
    """
    # Create the hyperlink relationship
    r_id = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(ns.qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Set Calibri font, 13.5 size, underlined, and blue color (Accent 1)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(ns.qn("w:ascii"), "Calibri")
    rFonts.set(ns.qn("w:hAnsi"), "Calibri")

    sz = OxmlElement("w:sz")
    sz.set(ns.qn("w:val"), "27")  # 13.5 pt (size * 2)

    u = OxmlElement("w:u")
    u.set(ns.qn("w:val"), "single")  # Underlined

    color = OxmlElement("w:color")
    color.set(ns.qn("w:val"), "0000FF")  # Blue (Accent 1)

    rPr.append(rFonts)
    rPr.append(sz)
    rPr.append(u)
    rPr.append(color)
    r.append(rPr)

    # Add text
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)

    hyperlink.append(r)
    paragraph._element.append(hyperlink)

# Add articles to Word document
for index, row in df.iterrows():
    title = row["Title"]
    url = row["URL"]

    # Add hyperlink title with custom formatting
    p = doc.add_paragraph()
    add_hyperlink(p, title, url)

    # Add URL below the title
    doc.add_paragraph(url)

# Save the document
doc.save(word_filename)

# **Save the document in the output_files folder**
word_file_path = "output_files/MEI.docx"
doc.save(word_file_path)

print(f"Word document created: {word_filename}")
# ----------------------------------
