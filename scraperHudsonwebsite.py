from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import pandas as pd
import time
from docx import Document  # For Word document creation
from docx.shared import Pt
from docx.oxml import OxmlElement, ns


# Initialize WebDriver
options = webdriver.ChromeOptions()
options.add_argument("--headless")  # Run in headless mode
driver = webdriver.Chrome(options=options)

# URL of Hudson Institute's research page
url = "https://www.hudson.org/search?hud-content-type=259&expert=&date-from=&date-to=&keywords=&topics=All&region=All"

# Open the webpage
driver.get(url)
time.sleep(5)  # Allow time for page to load

articles = []

try:
    # Wait for articles to load
    WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.CLASS_NAME, "c-horizontal-card__title"))
    )

    # Locate all articles
    article_elements = driver.find_elements(By.CLASS_NAME, "c-horizontal-card__title")

    for article in article_elements[:20]:  # Limit to 20 articles
        try:
            title_element = article.find_element(By.TAG_NAME, "span")
            title = title_element.text.strip()

            link = article.get_attribute("href")
            if not link.startswith("http"):
                link = "https://www.hudson.org" + link  # Ensure full URL

            articles.append({"Think Tank": "Hudson Institute", "Title": title, "URL": link})

        except Exception as e:
            print(f"⚠️ Skipping an article due to an error: {e}")
    
except Exception as e:
    print(f"❌ Error loading articles: {e}")

# Close the driver
driver.quit()

# Save results to CSV
df = pd.DataFrame(articles)
csv_filename = "Hudson_Articles.csv"
df.to_csv(csv_filename, index=False)


# **EXTENSION: Create a Word Document from CSV with Hyperlinked Titles**
# ----------------------------------
word_filename = "Hudson hyperlinks.docx"

# Read data from CSV
df = pd.read_csv("Hudson_articles.csv")

# Create a Word document
doc = Document()

# Add a title
doc.add_heading("Hudson Institute", level=1)

# Function to add a real clickable hyperlink with custom formatting
def add_hyperlink(paragraph, text, url):
    """
    Adds a real clickable hyperlink to a Word document with specific formatting.
    """
    # Create the hyperlink relationship
    r_id = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(ns.qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Set Calibri font, 13.5 size, underlined, and blue color (Accent 1)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(ns.qn("w:ascii"), "Calibri")
    rFonts.set(ns.qn("w:hAnsi"), "Calibri")

    sz = OxmlElement("w:sz")
    sz.set(ns.qn("w:val"), "27")  # 13.5 pt (size * 2)

    u = OxmlElement("w:u")
    u.set(ns.qn("w:val"), "single")  # Underlined

    color = OxmlElement("w:color")
    color.set(ns.qn("w:val"), "0000FF")  # Blue (Accent 1)

    rPr.append(rFonts)
    rPr.append(sz)
    rPr.append(u)
    rPr.append(color)
    r.append(rPr)

    # Add text
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)

    hyperlink.append(r)
    paragraph._element.append(hyperlink)

# Add articles to Word document
for index, row in df.iterrows():
    title = row["Title"]
    url = row["URL"]

    # Add hyperlink title with custom formatting
    p = doc.add_paragraph()
    add_hyperlink(p, title, url)

    # Add URL below the title
    doc.add_paragraph(url)

# Save the document
doc.save(word_filename)

# **Save the document in the output_files folder**
word_file_path = "output_files/Hudson Institute.docx"
doc.save(word_file_path)

print(f"✅ Word document created: {word_filename}")
# ----------------------------------


print(f"✅ Scraped {len(articles)} articles from Hudson Institute. Data saved to {csv_filename}.")
