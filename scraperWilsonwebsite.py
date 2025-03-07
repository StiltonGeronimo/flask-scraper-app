import time
import pandas as pd
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document  # For Word document creation
from docx.shared import Pt
from docx.oxml import OxmlElement, ns

# **Wilson Center Research Page URL**
wilson_url = "https://www.wilsoncenter.org/insight-analysis"

# **Configure Selenium WebDriver**
options = Options()
options.headless = False  # Set to False to see browser actions
options.add_argument("--no-sandbox")
options.add_argument("--disable-dev-shm-usage")

# Start Selenium WebDriver
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)

articles = []  # Store all articles

# **Function to scrape a page**
def scrape_page():
    global articles
    soup = BeautifulSoup(driver.page_source, "html.parser")

    for article in soup.find_all("h2", class_="title h4 -blue-600"):
        a_tag = article.find("a")
        if a_tag:
            title = a_tag.get_text(strip=True)  # Extract title
            link = urljoin(wilson_url, a_tag["href"])  # Ensure full URL
            articles.append({"Think Tank": "Wilson Center", "Title": title, "URL": link})

# **Step 1: Scrape the first page**
driver.get(wilson_url)
time.sleep(5)  # Allow JavaScript to load
scrape_page()

# **Step 2: Navigate to Page 2 & scrape**
try:
    page_2_button = driver.find_element(By.XPATH, '//a[@data-value="2"]')  # Find the Page 2 button
    driver.execute_script("arguments[0].click();", page_2_button)  # Click the button
    time.sleep(5)  # Wait for the new page to load
    scrape_page()  # Scrape the second page
    print("✅ Scraped Page 2")
except Exception as e:
    print(f"⚠️ Page 2 button not found: {e}. Skipping.")

# **Close browser**
driver.quit()

# **Save results to CSV**
if articles:
    df = pd.DataFrame(articles)
    df.to_csv("Wilson_Center_articles.csv", index=False)
    print("✅ Wilson Center articles scraped and saved to Wilson_Center_articles.csv")
else:
    print("⚠️ No articles found on Wilson Center's page.")

# **EXTENSION: Create a Word Document from CSV with Hyperlinked Titles**
# ----------------------------------
word_filename = "Wilson hyperlinks.docx"

# Read data from CSV
df = pd.read_csv("Wilson_Center_articles.csv")

# Create a Word document
doc = Document()

# Add a title
doc.add_heading("Wilson Center", level=1)

# Function to add a real clickable hyperlink with custom formatting
def add_hyperlink(paragraph, text, url):
    """
    Adds a real clickable hyperlink to a Word document with specific formatting.
    """
    # Create the hyperlink relationship
    r_id = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(ns.qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Set Calibri font, 13.5 size, underlined, and blue color (Accent 1)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(ns.qn("w:ascii"), "Calibri")
    rFonts.set(ns.qn("w:hAnsi"), "Calibri")

    sz = OxmlElement("w:sz")
    sz.set(ns.qn("w:val"), "27")  # 13.5 pt (size * 2)

    u = OxmlElement("w:u")
    u.set(ns.qn("w:val"), "single")  # Underlined

    color = OxmlElement("w:color")
    color.set(ns.qn("w:val"), "0000FF")  # Blue (Accent 1)

    rPr.append(rFonts)
    rPr.append(sz)
    rPr.append(u)
    rPr.append(color)
    r.append(rPr)

    # Add text
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)

    hyperlink.append(r)
    paragraph._element.append(hyperlink)

# Add articles to Word document
for index, row in df.iterrows():
    title = row["Title"]
    url = row["URL"]

    # Add hyperlink title with custom formatting
    p = doc.add_paragraph()
    add_hyperlink(p, title, url)

    # Add URL below the title
    doc.add_paragraph(url)

# Save the document
doc.save(word_filename)

# **Save the document in the output_files folder**
word_file_path = "output_files/Wilson Center.docx"
doc.save(word_file_path)

print(f"✅ Word document created: {word_filename}")
# ----------------------------------

