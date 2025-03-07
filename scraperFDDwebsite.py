import time
import pandas as pd
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document  # For Word document creation
from docx.shared import Pt
from docx.oxml import OxmlElement, ns

# **FDD Research Page URL**
fdd_url = "https://www.fdd.org/category/analysis/"

# **Configure Selenium WebDriver**
options = Options()
options.headless = False  # Set to False to see browser actions
options.add_argument("--no-sandbox")
options.add_argument("--disable-dev-shm-usage")

# Start Selenium WebDriver
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)

articles = []  # Store all scraped articles

# **Function to scrape a page**
def scrape_page():
    global articles
    soup = BeautifulSoup(driver.page_source, "html.parser")

    for article in soup.find_all("h4", class_="post-title"):
        title = article.get_text(strip=True)

        # Find the nearest <a> tag for the article link
        a_tag = article.find_parent("a", href=True)
        link = urljoin(fdd_url, a_tag["href"]) if a_tag else "No URL"

        articles.append({"Think Tank": "FDD", "Title": title, "URL": link})

# **Step 1: Scrape the first page**
driver.get(fdd_url)
time.sleep(5)  # Allow JavaScript to load
scrape_page()

# **Step 2: Click on Page 2 & Scrape**
try:
    page_2_button = driver.find_element(By.XPATH, '//span[@class="pagination-not-current" and text()="2"]')
    driver.execute_script("arguments[0].click();", page_2_button)  # Click the button
    time.sleep(5)  # Wait for the new page to load
    scrape_page()  # Scrape the second page
    print("✅ Scraped Page 2")
except Exception as e:
    print(f"⚠️ Could not find Page 2 button: {e}. Skipping.")

# **Step 3: Click on Page 3 & Scrape**
try:
    page_3_button = driver.find_element(By.XPATH, '//span[@class="pagination-not-current" and text()="3"]')
    driver.execute_script("arguments[0].click();", page_3_button)  # Click the button
    time.sleep(5)  # Wait for the new page to load
    scrape_page()  # Scrape the third page
    print("✅ Scraped Page 3")
except Exception as e:
    print(f"⚠️ Could not find Page 3 button: {e}. Skipping.")

# **Close browser**
driver.quit()

# **Save results to CSV**
if articles:
    df = pd.DataFrame(articles)
    df.to_csv("FDD_articles.csv", index=False)
    print("✅ FDD articles scraped and saved to FDD_articles.csv")
else:
    print("⚠️ No articles found on FDD's page.")

    
# **EXTENSION: Create a Word Document from CSV with Hyperlinked Titles**
# ----------------------------------
word_filename = "FDD hyperlinks.docx"

# Read data from CSV
df = pd.read_csv("FDD_articles.csv")

# Create a Word document
doc = Document()

# Add a title
doc.add_heading("FDD", level=1)

# Function to add a real clickable hyperlink with custom formatting
def add_hyperlink(paragraph, text, url):
    """
    Adds a real clickable hyperlink to a Word document with specific formatting.
    """
    # Create the hyperlink relationship
    r_id = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(ns.qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Set Calibri font, 13.5 size, underlined, and blue color (Accent 1)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(ns.qn("w:ascii"), "Calibri")
    rFonts.set(ns.qn("w:hAnsi"), "Calibri")

    sz = OxmlElement("w:sz")
    sz.set(ns.qn("w:val"), "27")  # 13.5 pt (size * 2)

    u = OxmlElement("w:u")
    u.set(ns.qn("w:val"), "single")  # Underlined

    color = OxmlElement("w:color")
    color.set(ns.qn("w:val"), "0000FF")  # Blue (Accent 1)

    rPr.append(rFonts)
    rPr.append(sz)
    rPr.append(u)
    rPr.append(color)
    r.append(rPr)

    # Add text
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)

    hyperlink.append(r)
    paragraph._element.append(hyperlink)

# Add articles to Word document
for index, row in df.iterrows():
    title = row["Title"]
    url = row["URL"]

    # Add hyperlink title with custom formatting
    p = doc.add_paragraph()
    add_hyperlink(p, title, url)

    # Add URL below the title
    doc.add_paragraph(url)

# Save the document
doc.save(word_filename)

# **Save the document in the output_files folder**
word_file_path = "output_files/FDD.docx"
doc.save(word_file_path)

print(f"✅ Word document created: {word_filename}")
# ----------------------------------
